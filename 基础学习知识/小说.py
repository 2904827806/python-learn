import requests
from lxml import etree
from docx import Document
import re
doc1 = Document()#创建文档
paragraph = doc1.add_paragraph()
for j in range(7,10000):
    url = fr'https://www.bq90.cc/book/80351/{j}.html'
    header = {
        'Cookie': 'Hm_lvt_2254e433738fcc1aed329cee161b448a=1712461098; Hm_lpvt_2254e433738fcc1aed329cee161b448a=1712461412',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36 Edg/123.0.0.0'
    }
    req = requests.get(url=url, headers=header)
    html = etree.HTML(req.text)
    a = html.xpath('//*[@id="chaptercontent"]/text()')
    doc1.add_paragraph('第'+ str(j - 6)+'章')
    if j == 7:
        b = a[:-2]
        for i in b:
            doc1.add_paragraph(i)
            print(i)
        doc1.add_paragraph()
        if b == []:
            break
        else:
            continue
    else:
        b = a[1:-2]
        # print(a)
        for i in b:
            doc1.add_paragraph(i)
            print(i)

        if b == []:
            break
        else:
            continue

doc1.save('小说.docx')

